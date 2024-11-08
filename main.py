from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
import os
import docx
from waitress import serve

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'docx'}

# Ensure the upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def clear_document(doc):
    for para in doc.paragraphs:
        doc.element.body.remove(para._element)
    for table in doc.tables:
        doc.element.body.remove(table._element)

def set_different_first_page(doc, value):
    for section in doc.sections:
        section.footer.distance_from_bottom = Inches(1)
        section.header.distance_from_top = Inches(1)
        section.different_first_page_header_footer = value

def normalize_text(text):
    """Normalize text to ignore case but retain formatting for readability."""
    return text.strip().lower()

def get_normalized_header(table):
    """Get a normalized header row from a table without altering spacing and punctuation."""
    return tuple(normalize_text(cell.text) for cell in table.rows[0].cells)

def merge_documents(files):
    merged_content = {}
    current_heading = ' '
    merged_content[current_heading] = {}

    for file in files:
        doc = Document(file)
        current_subheading = None
        content_order = [] 

        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = docx.text.paragraph.Paragraph(element, doc)
                if is_heading(para):
                    current_heading = para.text.strip()
                    current_subheading = None
                    if current_heading not in merged_content:
                        merged_content[current_heading] = {}

                elif is_subheading(para):
                    current_subheading = para.text.strip()
                    if current_subheading not in merged_content[current_heading]:
                        merged_content[current_heading][current_subheading] = []
                    content_order = merged_content[current_heading][current_subheading]

                elif current_heading:
                    if current_subheading:
                        content_order = merged_content[current_heading][current_subheading]
                    else:
                        content_order = merged_content[current_heading].setdefault(None, [])

                    content_order.append(('paragraph', para))

            elif element.tag.endswith('tbl'):
                table = docx.table.Table(element, doc)
                if current_heading:
                    if current_subheading:
                        content_order = merged_content[current_heading][current_subheading]
                    else:
                        content_order = merged_content[current_heading].setdefault(None, [])

                    content_order.append(('table', table))

    merged_doc = Document(files[0])
    clear_document(merged_doc)
    set_different_first_page(merged_doc, False)

    for heading, subcontent in merged_content.items():
        heading_paragraph = merged_doc.add_paragraph(heading)
        heading_run = heading_paragraph.runs[0]
        heading_run.bold = True
        heading_run.font.size = docx.shared.Pt(16)

        for subheading, elements in subcontent.items():
            if subheading:
                subheading_paragraph = merged_doc.add_paragraph(subheading)
                subheading_run = subheading_paragraph.runs[0]
                subheading_run.bold = True
                subheading_run.font.size = docx.shared.Pt(14)

            elements = merge_similar_tables(elements)

            for element_type, element in elements:
                if element_type == 'paragraph':
                    copy_paragraph_and_images(element, merged_doc)
                elif element_type == 'table':
                    copy_table(element, merged_doc)
                    merged_doc.add_paragraph("\n")

    output_path = 'merged_report.docx'
    merged_doc.save(output_path)
    return output_path

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        files = request.files.getlist('files')
        if files and all(allowed_file(f.filename) for f in files):
            file_paths = []
            for f in files:
                filename = secure_filename(f.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                f.save(file_path)
                file_paths.append(file_path)

            merged_file_path = merge_documents(file_paths)
            return send_file(merged_file_path, as_attachment=True)
    
    return render_template('upload.html')

def is_heading(para):
    return para.style.name.startswith('Heading')

def is_subheading(para):
    if para.style.name.startswith('Heading') and not para.style.name.startswith('Heading 1'):
        return True
    if para.runs and (para.runs[0].bold or para.runs[0].italic or para.runs[0].underline):
        return True
    return False

def copy_paragraph_and_images(source_para, target_doc):
    target_para = target_doc.add_paragraph()
    for run in source_para.runs:
        new_run = target_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        
        for drawing in run._element.xpath('.//w:drawing'):
            blip = drawing.xpath('.//a:blip')
            if blip:
                image_rId = blip[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                image_part = run.part.related_parts[image_rId]
                image_path = f"temp_image_{image_rId}.png"
                with open(image_path, "wb") as img_file:
                    img_file.write(image_part.blob)
                target_doc.add_picture(image_path, width=Inches(2))
                os.remove(image_path)

def set_cell_borders(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = docx.oxml.parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                    r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                                    r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                                    r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                                    r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                                    r'</w:tcBorders>')
    tcPr.append(tcBorders)

def copy_table(source_table, target_doc):
    new_table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    for row_idx, row in enumerate(source_table.rows):
        for col_idx, cell in enumerate(row.cells):
            new_table.cell(row_idx, col_idx).text = cell.text
            set_cell_borders(new_table.cell(row_idx, col_idx))

def merge_similar_tables(elements):
    combined_elements = []
    table_groups = {}

    for elem_type, table in elements:
        if elem_type == 'table':
            # Get the header of the table (normalize for easier comparison)
            normalized_header = get_normalized_header(table)
            first_column_value = table.cell(0, 0).text.strip()  # Value of the first column in the first row

            # Group tables by their normalized header, first column, and number of columns
            table_key = (normalized_header, len(table.columns), first_column_value)

            if table_key not in table_groups:
                table_groups[table_key] = []

            table_groups[table_key].append(table)
        else:
            combined_elements.append((elem_type, table))

    for (header, col_count, first_column_value), tables in table_groups.items():
        if len(tables) == 1:
            combined_elements.append(('table', tables[0]))
        else:
            # Merge tables that have the same header, same number of columns, and same first column value
            merged_table = Document().add_table(rows=0, cols=len(tables[0].columns))

            # Copy the header from the first table
            header_row = merged_table.add_row().cells
            for idx, cell in enumerate(tables[0].rows[0].cells):
                header_row[idx].text = cell.text

            # Append rows from all tables, skipping the header row from each table
            for table in tables:
                for row_idx, row in enumerate(table.rows[1:]):  # Skip the header row
                    # Ensure the first column matches and is added correctly
                    new_row = merged_table.add_row().cells
                    for col_idx, cell in enumerate(row.cells):
                        new_row[col_idx].text = cell.text

            combined_elements.append(('table', merged_table))

    return combined_elements

if __name__ == '__main__':
    app.run(debug=True)
